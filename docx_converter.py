from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

from .latex_utils import escape_latex, safe_filename
from .models import ConversionOptions, ConversionResult
from .tex_writer import wrap_document

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


def _style_name(p) -> str:
    try:
        return (p.style.name or "").lower()
    except Exception:
        return ""


def _heading_level(p) -> int | None:
    name = _style_name(p)
    m = re.search(r"(?:heading|titre)\s*(\d+)", name)
    if m:
        return min(max(int(m.group(1)), 1), 4)
    txt = (p.text or "").strip()
    if txt and len(txt) < 100 and txt.isupper() and len(txt.split()) <= 10:
        return 2
    return None


def _run_to_tex(run, clean_text: bool = True) -> str:
    text = escape_latex(run.text or "", clean_text=clean_text)
    if not text:
        return ""
    if run.bold:
        text = r"\textbf{" + text + "}"
    if run.italic:
        text = r"\textit{" + text + "}"
    if run.underline:
        text = r"\underline{" + text + "}"
    if run.font.superscript:
        text = r"\textsuperscript{" + text + "}"
    if run.font.subscript:
        text = r"\textsubscript{" + text + "}"
    if run.font.color and run.font.color.rgb:
        rgb = str(run.font.color.rgb)
        text = rf"{{\color[HTML]{{{rgb}}}{text}}}"
    return text


def _paragraph_runs(p, clean_text: bool = True, preserve_page_breaks: bool = True) -> str:
    chunks: list[str] = []
    for run in p.runs:
        chunks.append(_run_to_tex(run, clean_text=clean_text))
        for br in run._r.xpath('.//w:br'):
            typ = br.get(qn('w:type'))
            if typ == 'page':
                chunks.append(r"\newpage" if preserve_page_breaks else r"\\")
            elif typ == 'column':
                chunks.append(r"\columnbreak")
            else:
                chunks.append(r"\\")
        if run._r.xpath('.//w:tab'):
            chunks.append(r"\hspace*{1em}")
    return "".join(chunks) or escape_latex(p.text, clean_text=clean_text)


def _alignment(p) -> str:
    mapping = {WD_ALIGN_PARAGRAPH.CENTER: "center", WD_ALIGN_PARAGRAPH.RIGHT: "flushright", WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"}
    return mapping.get(p.alignment, "left")


def _list_info(p):
    ilvl = p._p.xpath('./w:pPr/w:numPr/w:ilvl')
    numid = p._p.xpath('./w:pPr/w:numPr/w:numId')
    if ilvl and numid:
        return max(0, int(ilvl[0].get(qn('w:val')) or 0)), int(numid[0].get(qn('w:val')) or 0)
    name = _style_name(p)
    if 'list bullet' in name:
        return 0, 1
    if 'list number' in name:
        return 0, 2
    return None


def _paragraph_tex(p, image_map: dict[str, str], options: ConversionOptions) -> tuple[str, int]:
    heading = _heading_level(p)
    text = _paragraph_runs(p, clean_text=options.clean_text, preserve_page_breaks=options.preserve_page_breaks).strip()
    image_count = 0
    image_names: list[tuple[str, float]] = []
    for blip in p._p.findall('.//a:blip', NS):
        rid = blip.get(qn('r:embed'))
        if rid and rid in image_map:
            width_cm = 8.0
            cur = blip
            extent = None
            for _ in range(8):
                cur = cur.getparent() if cur is not None else None
                if cur is None:
                    break
                extent = cur.find('wp:extent', NS)
                if extent is not None:
                    break
            if extent is not None:
                try:
                    width_cm = float(extent.get('cx')) / 360000.0
                except (TypeError, ValueError):
                    pass
            image_names.append((image_map[rid], width_cm))
    if image_names and options.preserve_images:
        for fname, width_cm in image_names:
            width = max(1.0, min(width_cm, 16.5))
            text += rf"\begin{{figure}}[H]\centering\includegraphics[width={width:.2f}cm]{{images/{fname}}}\end{{figure}}"
            image_count += 1
    if not text:
        return "", image_count
    if heading:
        cmd = {1: "section", 2: "subsection", 3: "subsubsection", 4: "paragraph"}[heading]
        return f"\\{cmd}{{{text}}}\n\n", image_count
    align = _alignment(p)
    if align == "center":
        return f"\\begin{{center}}{text}\\end{{center}}\n\n", image_count
    if align == "flushright":
        return f"\\begin{{flushright}}{text}\\end{{flushright}}\n\n", image_count
    if align == "justify":
        return f"\\begin{{justify}}{text}\\end{{justify}}\n\n", image_count
    return text + "\n\n", image_count


def _cell_text(cell, options: ConversionOptions) -> str:
    parts = []
    for p in cell.paragraphs:
        txt = _paragraph_runs(p, clean_text=options.clean_text, preserve_page_breaks=options.preserve_page_breaks).strip()
        if txt:
            parts.append(txt)
    return r"\newline{}".join(parts)


def _table_tex(table, index: int, options: ConversionOptions) -> str:
    rows = table.rows
    cols = len(rows[0].cells) if rows else 0
    if not cols:
        return ""
    env = "longtable" if options.optimize_tables and len(rows) >= 25 else "tabularx"
    if not options.optimize_tables:
        colspec = "|" + "|".join(["l"] * cols) + "|"
        lines = [r"\begin{table}[H]", r"\centering", rf"\begin{{tabular}}{{{colspec}}}", r"\hline"]
        for row in rows:
            vals = [_cell_text(cell, options).replace("\n", " ") for cell in row.cells]
            while len(vals) < cols:
                vals.append("")
            lines.append(" & ".join(vals[:cols]) + r" \\ \hline")
        lines += [r"\end{tabular}", rf"\caption{{Tableau {index}}}", r"\end{table}", ""]
        return "\n".join(lines)
    if env == "tabularx":
        colspec = "|" + "|".join([r">{\raggedright\arraybackslash}X"] * cols) + "|"
        lines = [r"\begin{table}[H]", r"\centering", rf"\begin{{tabularx}}{{\textwidth}}{{{colspec}}}", r"\hline"]
    else:
        width = max(2.0, round(15.8 / cols, 2))
        colspec = "|" + "|".join([rf">{{\raggedright\arraybackslash}}p{{{width:.2f}cm}}"] * cols) + "|"
        lines = [rf"\begin{{longtable}}{{{colspec}}}", rf"\caption{{Tableau {index}}}\\", r"\hline"]
    for row in rows:
        vals = [_cell_text(cell, options).replace("\n", " ") for cell in row.cells]
        while len(vals) < cols:
            vals.append("")
        lines.append(" & ".join(vals[:cols]) + r" \\ \hline")
    if env == "tabularx":
        lines += [r"\end{tabularx}", rf"\caption{{Tableau {index}}}", r"\end{table}", ""]
    else:
        lines += [r"\end{longtable}", ""]
    return "\n".join(lines)


def _extract_images(docx_path: Path, project_dir: Path) -> tuple[dict[str, str], int]:
    image_dir = project_dir / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    image_map: dict[str, str] = {}
    count = 0
    with zipfile.ZipFile(docx_path) as z:
        rels = ET.fromstring(z.read("word/_rels/document.xml.rels"))
        target_by_id = {}
        for rel in rels:
            rid = rel.attrib.get("Id")
            target = rel.attrib.get("Target", "")
            if rid and target.startswith("media/"):
                target_by_id[rid] = "word/" + target
        for rid, member in target_by_id.items():
            data = z.read(member)
            ext = Path(member).suffix.lower() or ".bin"
            name = f"image_{count + 1:03d}{ext}"
            (image_dir / name).write_bytes(data)
            image_map[rid] = name
            count += 1
    return image_map, count


def _header_footer_text(section, which: str) -> str | None:
    obj = section.header if which == "header" else section.footer
    vals = [_paragraph_runs(p).strip() for p in obj.paragraphs if p.text.strip()]
    return " \\quad ".join(vals) if vals else None


def convert_docx(source: Path, options: ConversionOptions) -> ConversionResult:
    project = options.output_dir / f"{safe_filename(source.stem)}_latex"
    project.mkdir(parents=True, exist_ok=True)
    warnings: list[str] = []
    try:
        doc = Document(source)
        image_map, image_count = _extract_images(source, project) if options.preserve_images else ({}, 0)
        body_parts: list[str] = []
        table_count = 0
        list_stack: list[str] = []
        paragraphs_by_element = {p._p: p for p in doc.paragraphs}
        tables_by_element = {t._tbl: t for t in doc.tables}
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                para = paragraphs_by_element.get(child)
                if para is not None:
                    info = _list_info(para)
                    if info:
                        # Close/open a simple list environment as nesting changes.
                        level, _ = info
                        env = "enumerate" if info[1] % 2 == 0 else "itemize"
                        while len(list_stack) > level + 1:
                            body_parts.append(r"\end{" + list_stack.pop() + "}\n")
                        if len(list_stack) < level + 1:
                            body_parts.append(rf"\begin{{{env}}}[leftmargin=*]" + "\n")
                            list_stack.append(env)
                        body_parts.append(r"\item " + _paragraph_runs(para, clean_text=options.clean_text, preserve_page_breaks=options.preserve_page_breaks).strip() + "\n")
                    else:
                        while list_stack:
                            body_parts.append(r"\end{" + list_stack.pop() + "}\n")
                        tex, _ = _paragraph_tex(para, image_map, options)
                        body_parts.append(tex)
            elif child.tag == qn("w:tbl"):
                while list_stack:
                    body_parts.append(r"\end{" + list_stack.pop() + "}\n")
                table = tables_by_element.get(child)
                if table is not None:
                    table_count += 1
                    body_parts.append(_table_tex(table, table_count, options))
        while list_stack:
            body_parts.append(r"\end{" + list_stack.pop() + "}\n")

        header = _header_footer_text(doc.sections[0], "header") if doc.sections else None
        footer = _header_footer_text(doc.sections[0], "footer") if doc.sections else None
        if header or footer:
            warnings.append("En-tête/pied de page simple repris ; les variantes complexes par section peuvent nécessiter une retouche.")
        tex = wrap_document("".join(body_parts), header=header, footer=footer)
        tex_path = project / "main.tex"
        tex_path.write_text(tex, encoding="utf-8")
        (project / "README.txt").write_text(
            "Projet généré par Docu2TeX. Compilez main.tex avec XeLaTeX (recommandé) ou LuaLaTeX.\n",
            encoding="utf-8",
        )
        return ConversionResult(True, source, project, tex_path, None, None, None, image_count, table_count, warnings, [], "")
    except Exception as exc:
        return ConversionResult(False, source, project, None, None, None, None, 0, 0, warnings, [f"{type(exc).__name__}: {exc}"], "")
