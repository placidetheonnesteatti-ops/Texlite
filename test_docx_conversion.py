from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from core.docx_converter import convert_docx
from core.models import ConversionOptions


def test_docx_to_tex(tmp_path: Path):
    src = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Mon titre", level=1)
    p = doc.add_paragraph()
    r = p.add_run("50% & test")
    r.bold = True
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nom"
    table.cell(0, 1).text = "Valeur"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"
    doc.save(src)
    result = convert_docx(src, ConversionOptions(output_dir=tmp_path))
    assert result.success
    text = result.tex_path.read_text(encoding="utf-8")
    assert r"\section{Mon titre}" in text
    assert r"50\% \& test" in text
    assert "tabularx" in text


def test_docx_complex_features(tmp_path: Path):
    src = tmp_path / "complex.docx"
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Document de test"
    section.footer.paragraphs[0].text = "Université"
    doc.add_heading("Titre principal", level=1)
    p = doc.add_paragraph("Texte normal avec ")
    r = p.add_run("gras")
    r.bold = True
    p.add_run(" et ")
    r = p.add_run("italique")
    r.italic = True
    for item in ["Premier point", "Deuxième point"]:
        doc.add_paragraph(item, style="List Bullet")
    table = doc.add_table(rows=3, cols=3)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cellule {i}-{j} & 50%"
    doc.save(src)
    result = convert_docx(src, ConversionOptions(output_dir=tmp_path))
    assert result.success
    text = result.tex_path.read_text(encoding="utf-8")
    assert "\\lhead{Document de test}" in text
    assert "\\rfoot{Université}" in text
    assert "\\begin{itemize}" in text
    assert "tabularx" in text
    assert "50\\%" in text


def test_french_heading_level_four_and_optional_page_breaks(tmp_path: Path):
    src = tmp_path / "features.docx"
    doc = Document()
    doc.styles["Heading 1"].name = "Titre 1"
    doc.add_paragraph("Titre français", style="Titre 1")
    doc.add_heading("Titre 4", level=4)
    paragraph = doc.add_paragraph("Avant")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("Après")
    doc.save(src)

    result = convert_docx(src, ConversionOptions(output_dir=tmp_path, preserve_page_breaks=False))
    assert result.success
    text = result.tex_path.read_text(encoding="utf-8")
    assert r"\section{Titre français}" in text
    assert r"\paragraph{Titre 4}" in text
    assert r"\newpage" not in text
